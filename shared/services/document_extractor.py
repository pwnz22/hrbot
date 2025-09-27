import os
from typing import Optional, Tuple
import PyPDF2
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocumentTextExtractor:
    """Service for extracting text from PDF and DOCX files"""

    @staticmethod
    def extract_text_from_file(file_path: str) -> Optional[str]:
        """
        Extract text from PDF or DOCX file

        Args:
            file_path: Path to the file

        Returns:
            Extracted text or None if extraction failed
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None

        file_extension = os.path.splitext(file_path)[1].lower()

        try:
            if file_extension == '.pdf':
                return DocumentTextExtractor._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return DocumentTextExtractor._extract_from_docx(file_path)
            else:
                logger.warning(f"Unsupported file format: {file_extension}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        return text.strip()

    @staticmethod
    def is_supported_format(filename: str) -> bool:
        """Check if file format is supported"""
        if not filename:
            return False
        extension = os.path.splitext(filename)[1].lower()
        return extension in ['.pdf', '.docx']